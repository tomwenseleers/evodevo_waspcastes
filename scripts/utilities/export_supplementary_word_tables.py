"""Export the final supplementary Word tables as UTF-8 TSV files.

This utility is retained for provenance. It is not required to run the R
analyses because the exported tables are versioned in output/supplemental_tables.
"""

from __future__ import annotations

import csv
import sys
from pathlib import Path

from docx import Document


TABLE_NAMES = (
    "TableS1",
    "TableS2",
    "TableS3",
    "TableS4",
    "TableS5A",
    "TableS5B",
    "TableS6",
    "TableS7",
    "TableS8",
    "TableS9",
    "TableS10",
    "TableS11",
    "TableS12",
    "TableS13",
)


def clean_cell(text: str) -> str:
    return " ".join(text.replace("\t", " ").replace("\r", " ").splitlines()).strip()


def main(source: Path, destination: Path) -> None:
    document = Document(source)
    if len(document.tables) != len(TABLE_NAMES):
        raise RuntimeError(
            f"Expected {len(TABLE_NAMES)} tables but found {len(document.tables)} in {source}"
        )

    destination.mkdir(parents=True, exist_ok=True)
    for name, table in zip(TABLE_NAMES, document.tables, strict=True):
        output = destination / f"{name}.tsv"
        with output.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle, delimiter="\t", lineterminator="\n")
            for row in table.rows:
                writer.writerow([clean_cell(cell.text) for cell in row.cells])
        print(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit(
            "Usage: export_supplementary_word_tables.py SOURCE.docx OUTPUT_DIRECTORY"
        )
    main(Path(sys.argv[1]), Path(sys.argv[2]))
